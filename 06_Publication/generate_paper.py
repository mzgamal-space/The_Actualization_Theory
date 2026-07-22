"""
generate_paper.py — Academic Publication Compiler (arXiv / IEEE style)
=======================================================================
Author: Mohamed Gamal Eldin Abdelaziz Noureldin
        Independent Researcher
        ORCID: 0009-0006-3991-1153
        Contact: mz.gamal@gmail.com

Generates:
  1. Actualizer_Engine_Paper.docx — full academic preprint embedding all 6 figures.
  2. Actualizer_Engine_Paper.md   — arXiv-formatted markdown manuscript.

Run:  python generate_paper.py
"""
import os, sys
import docx
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE    = os.path.dirname(os.path.abspath(__file__))
VIZ_DIR = os.path.join(HERE, '..', '04_Visualizations')
OUT_DOCX = os.path.join(HERE, 'Actualizer_Engine_Paper.docx')
OUT_MD   = os.path.join(HERE, 'Actualizer_Engine_Paper.md')

FIGURES = {
    "fig1": os.path.join(VIZ_DIR, 'fig1_hallucination_comparison.png'),
    "fig2": os.path.join(VIZ_DIR, 'fig2_repetition_suppression.png'),
    "fig3": os.path.join(VIZ_DIR, 'fig3_speed_comparison.png'),
    "fig4": os.path.join(VIZ_DIR, 'fig4_search_space_scaling.png'),
    "fig5": os.path.join(VIZ_DIR, 'fig5_v3u1_valuation_trajectory.png'),
    "fig6": os.path.join(VIZ_DIR, 'fig6_qca_parallel_speedup.png'),
}


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def build_docx(doc):
    for s in doc.sections:
        s.top_margin    = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin   = Inches(1.15)
        s.right_margin  = Inches(1.15)

    def p_normal(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 indent=0.0, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        if indent:
            p.paragraph_format.left_indent  = Inches(indent)
            p.paragraph_format.right_indent = Inches(indent)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.font.name   = 'Times New Roman'
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        return p

    def heading(text, level=1, num=""):
        sizes = {1: 13, 2: 11.5, 3: 11}
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(f"{num}  {text}" if num else text)
        r.font.name  = 'Times New Roman'
        r.font.size  = Pt(sizes.get(level, 11))
        r.font.bold  = True

    def equation(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(text)
        r.font.name   = 'Times New Roman'
        r.font.size   = Pt(11)
        r.font.italic = True

    def caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(12)
        r = p.add_run(text)
        r.font.name   = 'Times New Roman'
        r.font.size   = Pt(9.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    def figure(fig_key, cap_text):
        path = FIGURES.get(fig_key)
        if path and os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Inches(6.0))
        else:
            p_normal(f"[Figure not found: {fig_key}]", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        caption(cap_text)

    def add_table(headers, rows, header_bg="1F3864"):
        n_cols = len(headers)
        t = doc.add_table(rows=1 + len(rows), cols=n_cols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style     = 'Table Grid'
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            r = hdr[i].paragraphs[0].runs[0]
            r.font.name  = 'Times New Roman'
            r.font.bold  = True
            r.font.size  = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(hdr[i], header_bg)
        for ri, row_data in enumerate(rows):
            cells = t.rows[ri + 1].cells
            bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
            for ci, val in enumerate(row_data):
                cells[ci].text = str(val)
                r = cells[ci].paragraphs[0].runs[0]
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                _set_cell_bg(cells[ci], bg)
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(8)

    # ── TITLE & ABSTRACT ──
    p_normal(
        "The Actualizer Engine, FDSA, and QCA Parallel Architecture:\n"
        "A Unified Top-Down Steering & Clustered Substrate Framework for Factual Grounding "
        "in Large Language Models",
        bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=12, space_after=10
    )
    p_normal(
        "Mohamed Gamal Eldin Abdelaziz Noureldin\n"
        "Independent Researcher  |  ORCID: 0009-0006-3991-1153\n"
        "Contact: mz.gamal@gmail.com",
        size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14
    )

    p_normal("Abstract", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    p_normal(
        "Large language models operating under bottom-up Maximum Likelihood Estimation (MLE) suffer from "
        "combinatorial search explosion O(M^N), causing hallucination cascades, semantic drift, and repetition loops. "
        "We present a unified architecture combining the Actualizer Engine, the Fractal Deduction Search Algorithm (FDSA), "
        "and the Quench-Cluster Algorithm (QCA) Parallel Engine. The FDSA prunes invalid logit search space by up to 99.99% "
        "via isomorphic anchoring and dimensional truncation (D = ln V / ln(1/k_ref)). The QCA engine crystallizes un-clustered "
        "substrates into K independent sub-problems using the canonical RGG Quench Temperature T_q^RGG, reducing overall steering "
        "complexity from O(N^2) to O(N^2/K) across parallel worker processes or vectorized JAX units. The Actualizer Engine contractively "
        "steers residual distributions to zero-drift fixed points S_* via Banach contractive mappings (k = 0.45) guided by the five "
        "Conceptual Primes (Order, Justice, Mercy, Knowledge, Power). Empirical benchmarks at V = 1,000 to 100,000 demonstrate up to "
        "12.4x pre-inference sampling speedup, 3.69x dataset parallel speedup at N=200, and 100% factual grounding under strong distractor bait.",
        italic=True, indent=0.5, space_after=14
    )

    # ── SECTIONS ──
    heading("1. Introduction", 1)
    p_normal(
        "Modern autoregressive Transformer architectures achieve state-of-the-art performance through bottom-up "
        "statistical next-token prediction. However, they remain bounded by a fundamental epistemological pathology: "
        "in sparse-data or adversarial contexts, the probability distribution over the vocabulary 'smears' — all tokens "
        "receive nearly equal probability mass and the model becomes blind to structural validity constraints. "
        "This leads to hallucination cascades and repetition loops."
    )
    p_normal(
        "We present a structurally grounded solution: the FDSA + Actualizer + QCA Parallel unified framework. "
        "The FDSA enforces top-down dimensional truncation before inference, QCA crystallizes large datasets into K parallel "
        "sub-problems, and the Actualizer Engine contractively steers the residual substrate toward a zero-drift actualized state."
    )

    heading("2. Theoretical Foundation: The Conceptual Primes", 1)
    p_normal(
        "The Actualizer Engine evaluates candidate tokens against five invariant boundary metrics called the Conceptual Primes:"
    )
    add_table(
        ["Prime", "Role in Generation", "LLM Equivalent", "Drift on Violation"],
        [
            ["Order",    "Enforces local syntactic alignment",       "Grammar / syntax rules",        "Repetition cascade"],
            ["Justice",  "Balances global semantic distribution",    "Prompt boundary adherence",     "Topic drift"],
            ["Mercy",    "Decays local entropy overloads (Mercy = k)","Probability mass smoothing",    "Overconfidence collapse"],
            ["Knowledge","Projects downstream causal risk",          "Lookahead / future coherence",  "Sequence dead-end"],
            ["Power",    "Executes causal snap (bifurcation gated)", "Token selection (argmax)",      "Indecision / flat tie"],
        ]
    )
    caption("Table 1. The five Conceptual Primes and their roles in generation.")

    heading("3. Mathematical Framework", 1)
    heading("3.1  The Uncollapsed Probability Substrate", 2)
    equation("U_0 = softmax(z) = exp(z_v) / sum_v exp(z_v)  in R^V")

    heading("3.2  The Fractal Deduction Search Algorithm (FDSA)", 2)
    p_normal("(a) Isomorphic Anchoring: cosine matches Prime profile P(U) to reference domain P(R), inheriting k_ref.")
    equation("P(U) ~= P(R)  =>  extract  k_ref  from reference domain R")
    p_normal("(b) Actualization Fractal Dimension:")
    equation("D = ln(V) / ln(1 / k_ref)")
    p_normal("(c) Vectorized Logit Masking:")
    equation("z(v) = -inf  if  z(v) < -D * 1.5  OR  v not in grammar[last_token]")

    heading("3.3  The Drift Tensor & V3_U1 Structural Entropy", 2)
    p_normal("Systemic Structural Entropy H(R) uses the V3_U1-corrected squared magnitude defect term:")
    equation("H(R) = Var(alpha) + (sum_i alpha_i^2 - 1)^2")
    p_normal("Tripartite Drift Tensor D_mu_nu:")
    equation("D_mu_nu = w_L * D_local + w_G * D_global + w_F * D_future")

    heading("3.4  Vacuum Brake, Banach Contraction & Valuation Trajectory", 2)
    equation("U_braked(v) = U_n(v) * exp(-D(v) / tau)   (renormalised)")
    equation("U_{n+1} = k * U_braked + (1 - k) * U_n,   where k = 0.45  (Mercy = k)")
    equation("nu_t(A) = 1 - H(R_A(t)) / H_max   in [0, 1]")

    heading("3.5  Tr(D_mu_nu) Bifurcation Criterion & Causal Snap", 2)
    p_normal("Bifurcation Criterion (Theorem 3.3):")
    equation("Tr(D_mu_nu) <= tau_bifurcation  =>  S* = argmax U  (Actualization Branch)")
    equation("Tr(D_mu_nu) >  tau_bifurcation  =>  Dissolution Branch  (Fallback)")

    heading("3.6  Quench-Cluster Algorithm (QCA) Parallel Engine", 2)
    p_normal("To solve large dataset substrates of N nodes, QCA crystallizes N nodes into K independent clusters via RGG Quench Temperature:")
    equation("T_q^RGG = gamma * sqrt( A * ln(N/K) / (pi * N) )")
    p_normal("By CKT White Paper v3 §7.2 Theorem 2 Corollary, splitting N into K parallel clusters reduces steering work from O(N^2) to K * O((N/K)^2) = O(N^2/K) — a factor-K parallel acceleration.")

    heading("4. JAX & Multi-Process Execution Architecture", 1)
    p_normal("The QCAParallelEngine supports dual high-performance execution backends:")
    add_table(
        ["Backend", "Mechanism", "Hardware Optimization", "Target Workload"],
        [
            ["Processes", "Python ProcessPoolExecutor", "Multi-core CPU scaling", "General CPU environments"],
            ["JAX",       "Vectorized jnp.ndarray & @jax.jit", "GPU / TPU SIMD matrix units", "High-throughput inference"],
            ["Auto",      "Dynamic capability detection", "Automatic selection", "Production deployment"],
        ]
    )
    caption("Table 2. Execution backends supported by the QCA Parallel Engine.")

    heading("5. Experimental Benchmarks and Results", 1)

    heading("5.1  Hallucination Resistance", 2)
    p_normal("A strong distractor token (+8.0) was injected at every step. Baseline collapsed at step 2, while FDSA+Actualizer maintained 100% factual grounding over 30 steps.")
    figure("fig1", "Figure 1. Hallucination resistance comparison (V=500, 30 steps). Baseline collapses; FDSA+Actualizer maintains 100% grounding.")

    heading("5.2  Repetition Loop Suppression", 2)
    p_normal("Order Prime successfully suppressed repetition loops, increasing token diversity by 3.1x under adversarial +4.5 logit boosts.")
    figure("fig2", "Figure 2. Repetition suppression: repeat rate and token diversity comparison.")

    heading("5.3  Pre-Inference Speed Sweep (V = 1k to 100k)", 2)
    p_normal("FDSA pruner achieved up to 12.4x speedup at V=100,000 by pruning 99.99% of invalid logits before softmax.")
    figure("fig3", "Figure 3. Speed comparison, speedup factor, and pruning rate across V = 1k to 100k.")

    heading("5.4  Search Space Scaling Analysis (N = 4 to 18)", 2)
    p_normal("Combinatorial O(3^N) search space vs FDSA dimensional truncation O(N^D). At N=18, search space is reduced by >99.99%.")
    figure("fig4", "Figure 4. Asymptotic scaling: Combinatorial O(3^N) vs FDSA O(N^D) across N=4..18.")

    heading("5.5  V3_U1 Valuation Trajectory and Compliance", 2)
    p_normal("Valuation trajectory nu_t(A) and Tr(D_mu_nu) bifurcation threshold verifications across clean, moderate, and distractor substrates.")
    figure("fig5", "Figure 5. Valuation trajectory nu_t(A), drift tensor bifurcation, and V3_U1 compliance tests.")

    heading("5.6  QCA Parallel Engine Speedup Benchmark", 2)
    p_normal("Benchmark results for dataset sizes N = (20, 40, 80, 120, 200) with K=5 clusters and V=1,000 vocabulary size:")
    add_table(
        ["Dataset Size (N)", "Sequential Time (ms)", "Parallel QCA Time (ms)", "Speedup Factor", "Mean Valuation (nu)"],
        [
            ["N = 20",  "3,439.09 ms", "4,514.05 ms", "0.76x (Overhead bound)", "0.3792"],
            ["N = 40",  "6,760.25 ms", "4,931.68 ms", "1.37x",                 "0.4198"],
            ["N = 80",  "13,426.56 ms","7,325.66 ms", "1.83x",                 "0.5915"],
            ["N = 120", "20,164.83 ms","8,827.70 ms", "2.28x",                 "0.5408"],
            ["N = 200", "33,701.77 ms","9,142.00 ms", "3.69x (Approaching K=5x)", "0.2132"],
        ]
    )
    caption("Table 3. Empirical QCA Parallel Engine speedup across dataset sizes N=20..200 (K=5, V=1000).")
    figure("fig6", "Figure 6. QCA Parallel Engine acceleration: execution latency vs dataset size N and speedup factor.")

    heading("6. Conclusion", 1)
    p_normal(
        "We presented the unified Actualizer Engine, FDSA, and QCA Parallel framework. "
        "By combining pre-inference logit pruning, RGG quench-clustering, and Banach contractive steering, "
        "the architecture achieves 99.99% search space reduction, 12.4x pre-inference speedup, 3.69x dataset parallel speedup, "
        "and 100% factual grounding under adversarial conditions."
    )

    heading("References", 1)
    for ref in [
        "[1] Noureldin, M.G.E.A. — The Actualization Theory V3_U1 (2026). Independent Research.",
        "[2] Noureldin, M.G.E.A. — Quench Cluster Algorithm (QCA) & Parallel Actualizer (2026).",
        "[3] Vaswani et al. — Attention Is All You Need. NeurIPS 2017.",
        "[4] Banach, S. — Sur les operations dans les ensembles abstraits. Fund. Math. 1922.",
        "[5] Jax Development Team — JAX: Composable transformations of Python+NumPy (2018).",
    ]:
        p_normal(ref, size=10, space_after=3)

    doc.save(OUT_DOCX)
    print(f"DOCX Paper saved -> {OUT_DOCX}")


def build_markdown():
    md_content = """# The Actualizer Engine, FDSA, and QCA Parallel Architecture: A Unified Top-Down Steering & Clustered Substrate Framework for Factual Grounding in Large Language Models

**Mohamed Gamal Eldin Abdelaziz Noureldin**  
*Independent Researcher* | ORCID: [0009-0006-3991-1153](https://orcid.org/0009-0006-3991-1153) | Contact: `mz.gamal@gmail.com`

---

## Abstract

Large language models operating under bottom-up Maximum Likelihood Estimation (MLE) suffer from combinatorial search explosion $O(M^N)$, causing hallucination cascades, semantic drift, and repetition loops. We present a unified architecture combining the **Actualizer Engine**, the **Fractal Deduction Search Algorithm (FDSA)**, and the **Quench-Cluster Algorithm (QCA) Parallel Engine**. The FDSA prunes invalid logit search space by up to 99.99% via isomorphic anchoring and dimensional truncation ($D = \\ln V / \\ln(1/k_{ref})$). The QCA engine crystallizes un-clustered substrates into $K$ independent sub-problems using the canonical RGG Quench Temperature $T_q^{RGG}$, reducing overall steering complexity from $O(N^2)$ to $O(N^2/K)$ across parallel worker processes or vectorized JAX units. The Actualizer Engine contractively steers residual distributions to zero-drift fixed points $S^*$ via Banach contractive mappings ($k = 0.45$) guided by the five **Conceptual Primes** (Order, Justice, Mercy, Knowledge, Power). Empirical benchmarks at $V = 1,000$ to $100,000$ demonstrate up to 12.4× pre-inference sampling speedup, 3.69× dataset parallel speedup at $N=200$, and 100% factual grounding under strong distractor bait.

---

## 1. Introduction

Modern autoregressive Transformer architectures achieve state-of-the-art performance through bottom-up statistical next-token prediction. However, they remain bounded by a fundamental epistemological pathology: in sparse-data or adversarial contexts, the probability distribution over the vocabulary "smears" — all tokens receive nearly equal probability mass and the model becomes blind to structural validity constraints.

We present a structurally grounded solution: the FDSA + Actualizer + QCA Parallel unified framework. The FDSA enforces top-down dimensional truncation before inference, QCA crystallizes large datasets into $K$ parallel sub-problems, and the Actualizer Engine contractively steers the residual substrate toward a zero-drift actualized state.

---

## 2. Theoretical Foundation: The Conceptual Primes

The Actualizer Engine evaluates candidate tokens against five invariant boundary metrics called the Conceptual Primes:

| Prime | Role in Generation | LLM Equivalent | Drift on Violation |
|---|---|---|---|
| **Order** | Enforces local syntactic alignment | Grammar / syntax rules | Repetition cascade |
| **Justice** | Balances global semantic distribution | Prompt boundary adherence | Topic drift |
| **Mercy** | Decays local entropy overloads ($k$ IS Mercy) | Probability mass smoothing | Overconfidence collapse |
| **Knowledge** | Projects downstream causal risk | Lookahead / future coherence | Sequence dead-end |
| **Power** | Executes causal snap (bifurcation gated) | Token selection (argmax) | Indecision / flat tie |

---

## 3. Mathematical Framework

### 3.1 The Uncollapsed Probability Substrate
$$U_0 = \\text{softmax}(z) = \\frac{\\exp(z_v)}{\\sum_v \\exp(z_v)} \\in \\mathbb{R}^V$$

### 3.2 The Fractal Deduction Search Algorithm (FDSA)
1. **Isomorphic Anchoring:** $P(U) \\cong P(R) \\implies$ extract $k_{ref}$ from reference domain $R$.
2. **Actualization Fractal Dimension:**
   $$D = \\frac{\\ln(V)}{\\ln(1 / k_{ref})}$$
3. **Vectorized Logit Masking:**
   $$z(v) = -\\infty \\quad \\text{if } z(v) < -1.5 D \\quad \\text{or } v \\notin \\text{grammar}[\\text{last\\_token}]$$

### 3.3 The Drift Tensor & V3_U1 Structural Entropy
$$H(R) = \\text{Var}(\\alpha) + \\left(\\sum_i \\alpha_i^2 - 1\\right)^2$$
$$D_{\\mu\\nu} = w_L D_{\\text{local}} + w_G D_{\\text{global}} + w_F D_{\\text{future}}$$

### 3.4 Vacuum Brake & Banach Contraction
$$U_{\\text{braked}}(v) = U_n(v) \\exp\\left(-\\frac{D(v)}{\\tau}\\right)$$
$$U_{n+1} = k U_{\\text{braked}} + (1 - k) U_n, \\quad k = 0.45 \\quad (\\text{Mercy} = k)$$
$$\\nu_t(A) = 1 - \\frac{H(R_A(t))}{H_{\\max}} \\in [0, 1]$$

### 3.5 $\\text{Tr}(D_{\\mu\\nu})$ Bifurcation Criterion & Causal Snap
$$\\text{Tr}(D_{\\mu\\nu}) \\le \\tau_{\\text{bifurcation}} \\implies S^* = \\arg\\max U \\quad (\\text{Actualization Branch})$$
$$\\text{Tr}(D_{\\mu\\nu}) > \\tau_{\\text{bifurcation}} \\implies \\text{Dissolution Branch (Fallback)}$$

### 3.6 Quench-Cluster Algorithm (QCA) Parallel Engine
$$T_q^{\\text{RGG}} = \\gamma \\sqrt{\\frac{A \\ln(N/K)}{\\pi N}}$$
Splitting $N$ nodes into $K$ parallel clusters reduces steering work from $O(N^2)$ to $K \\cdot O((N/K)^2) = O(N^2/K)$ — a factor-$K$ parallel acceleration.

---

## 4. Execution Architecture (Processes & JAX)

| Backend | Mechanism | Hardware Optimization | Target Workload |
|---|---|---|---|
| **Processes** | Python `ProcessPoolExecutor` | Multi-core CPU scaling | General CPU environments |
| **JAX** | Vectorized `jnp.ndarray` & `@jax.jit` | GPU / TPU SIMD matrix units | High-throughput inference |
| **Auto** | Dynamic capability detection | Automatic selection | Production deployment |

---

## 5. Experimental Benchmarks and Visualizations

### 5.1 Hallucination Resistance
![Figure 1 — Hallucination Resistance](../04_Visualizations/fig1_hallucination_comparison.png)  
*Figure 1: Baseline immediately collapses under distractor bait (+8.0 logit); FDSA+Actualizer maintains 100% factual grounding over 30 steps.*

### 5.2 Repetition Suppression
![Figure 2 — Repetition Suppression](../04_Visualizations/fig2_repetition_suppression.png)  
*Figure 2: Order Prime repetition suppression increases token diversity by 3.1× under adversarial repetition bait.*

### 5.3 Pre-Inference Speed Sweep ($V = 1\\text{k} \\to 100\\text{k}$)
![Figure 3 — Speed Sweep](../04_Visualizations/fig3_speed_comparison.png)  
*Figure 3: FDSA pruner achieves up to 12.4× speedup at $V=100,000$ by pruning 99.99% of invalid logits before softmax.*

### 5.4 Search Space Scaling ($N = 4 \\to 18$)
![Figure 4 — Search Space Scaling](../04_Visualizations/fig4_search_space_scaling.png)  
*Figure 4: Asymptotic search space reduction: Combinatorial $O(3^N)$ vs FDSA $O(N^D)$. At $N=18$, search space is reduced by >99.99%.*

### 5.5 V3_U1 Valuation Trajectory & Compliance
![Figure 5 — Valuation Trajectory](../04_Visualizations/fig5_v3u1_valuation_trajectory.png)  
*Figure 5: Valuation trajectory $\\nu_t(A)$, drift tensor bifurcation threshold $\\tau=5.0$, and 6 targeted FIX compliance tests.*

### 5.6 QCA Parallel Engine Speedup Benchmark

| Dataset Size ($N$) | Sequential Time (ms) | Parallel QCA Time (ms) | Speedup Factor | Mean Valuation ($\\nu$) |
|---|---|---|---|---|
| $N = 20$ | 3,439.09 ms | 4,514.05 ms | 0.76× | 0.3792 |
| $N = 40$ | 6,760.25 ms | 4,931.68 ms | 1.37× | 0.4198 |
| $N = 80$ | 13,426.56 ms | 7,325.66 ms | 1.83× | 0.5915 |
| $N = 120$ | 20,164.83 ms | 8,827.70 ms | 2.28× | 0.5408 |
| $N = 200$ | 33,701.77 ms | 9,142.00 ms | **3.69×** (Approaching $K=5\\times$) | 0.2132 |

![Figure 6 — QCA Parallel Speedup](../04_Visualizations/fig6_qca_parallel_speedup.png)  
*Figure 6: QCA Parallel Engine acceleration: execution latency vs dataset size $N$ and empirical speedup factor up to 3.69× ($K=5, V=1,000$).*

---

## 6. Conclusion

We presented the unified Actualizer Engine, FDSA, and QCA Parallel framework. By combining pre-inference logit pruning, RGG quench-clustering, and Banach contractive steering, the architecture achieves 99.99% search space reduction, 12.4× pre-inference speedup, 3.69× dataset parallel speedup, and 100% factual grounding under adversarial conditions.

---

## References

1. Noureldin, M.G.E.A. — *The Actualization Theory V3_U1* (2026). Independent Research.
2. Noureldin, M.G.E.A. — *Quench Cluster Algorithm (QCA) & Parallel Actualizer* (2026).
3. Vaswani et al. — *Attention Is All You Need*. NeurIPS 2017.
4. Banach, S. — *Sur les operations dans les ensembles abstraits*. Fund. Math. 1922.
5. JAX Development Team — *JAX: Composable transformations of Python+NumPy* (2018).
"""
    with open(OUT_MD, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"Markdown Paper saved -> {OUT_MD}")


if __name__ == "__main__":
    doc = docx.Document()
    build_docx(doc)
    build_markdown()
